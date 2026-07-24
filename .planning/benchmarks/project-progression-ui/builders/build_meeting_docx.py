from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "fixtures" / "01_项目推进模拟周会纪要.docx"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x6F, 0x7B)
DOCUMENT_FONT = "Arial Unicode MS"


def set_font(run, size=11, color=RGBColor(0, 0, 0), bold=False, italic=False):
    run.font.name = DOCUMENT_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def apply_style_tokens(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = DOCUMENT_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = DOCUMENT_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_metadata(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    set_font(paragraph.add_run(f"{label}："), bold=True)
    set_font(paragraph.add_run(value))


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    set_font(paragraph.add_run(text))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, (label, width) in enumerate(zip(headers, widths)):
        cell = header.cells[index]
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(cell, "F2F4F7")
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(paragraph.add_run(label), size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, (value, width) in enumerate(zip(row, widths)):
            cell = cells[index]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in (0, 5) else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            set_font(paragraph.add_run(str(value)), size=9.3)
    return table


def build():
    doc = Document()
    apply_style_tokens(doc)
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("UI BENCHMARK · 项目推进模拟材料"), size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("仅用于 AI 项目作战管理平台 UI Benchmark，不代表真实项目事实"), size=8.5, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("项目推进模拟周会纪要"), size=23, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("第 2 次跨作战单元推进会 · Benchmark 样本"), size=14, color=RGBColor(55, 55, 55))

    add_metadata(doc, "会议日期", "2026-07-23 14:00-15:20")
    add_metadata(doc, "主持人", "联合指挥部 PMO（模拟）")
    add_metadata(doc, "参会范围", "研发、技术服务、产品、平台、财务、行政作战单元")
    add_metadata(doc, "材料性质", "合成测试材料；用于验证 UI 摄入、证据定位、冲突提示与更新提案")

    doc.add_heading("1. 本次会议确认的事实", level=1)
    add_bullet(doc, "技术服务作战单元负责的“XuguDB 技术知识库首轮”已完成 18/20 个知识主题，当前进度 90%，负责人冯治龙，计划在 2026-07-31 前完成首轮验收。")
    add_bullet(doc, "产品作战单元负责的“AI 效果台账”已完成模板首版，并录入 12 条试点记录；当前进度按会上口径暂记 60%，负责人王安迪。")
    add_bullet(doc, "平台作战单元的“平台作战单元正式规划”尚未补齐开始日期、结束日期和明确负责人，状态保持待确认，不得推断为已启动。")
    add_bullet(doc, "行政作战单元已完成 56 个技能模块安装，但算力分级调度仍是办公高频场景验证的前置阻塞。")

    doc.add_heading("2. 行动项与计划", level=1)
    add_table(
        doc,
        ["行动项", "作战单元", "负责人", "截止日期", "状态", "证据要求"],
        [
            ["公司级知识系统方案与治理规范：形成 v0.4 评审稿", "技术服务", "魏粤川", "2026-08-05", "进行中", "补充评审纪要与版本链接"],
            ["AI 效果台账：统一指标口径并补齐 20 条试点记录", "产品", "王安迪", "2026-08-12", "进行中", "上传指标表与抽样验收记录"],
            ["财务数据治理与安全分类：提交分类审批", "财务", "陈文斌", "2026-08-07", "阻塞", "需信息安全负责人签字"],
            ["平台作战单元正式规划：补齐负责人和排期", "平台", "曹茜", "2026-07-30", "待确认", "会议后书面确认，不得自动补造"],
        ],
        [1.80, 0.82, 0.66, 0.82, 0.70, 1.70],
    )

    doc.add_heading("3. 风险与阻塞", level=1)
    add_bullet(doc, "高风险：财务“基础核算与 OCR 试点验证”因数据分类审批未完成，预计最早 2026-08-12 才能进入真实票据验证。负责人陈文斌；缓解措施为先完成脱敏样本验证并提交安全评审。")
    add_bullet(doc, "中风险：AI 效果台账的完成度存在 60% 与群聊中“差不多八成”两种口径，本纪要以 60% 作为会议确认值，群聊口径仅作待核实线索。")

    doc.add_heading("4. 阶段成果", level=1)
    add_bullet(doc, "成果名称：公司级知识系统信息架构 v0.3。结果：联合评审通过，可作为下一版治理规范的结构基线。来源：本次会议材料与技术服务单元评审记录。完成日期：2026-07-22。")

    doc.add_heading("5. 明确不应自动处理的内容", level=1)
    add_bullet(doc, "群聊中的“基本搞定”“差不多八成”等表达不是正式完成状态，除非有明确负责人、日期和证据，不应直接更新项目。")
    add_bullet(doc, "任何负责人、任务日期、进度、风险关闭或成果归档建议都必须以结构化 ChangeProposal 呈现，经人工审核、合并到草稿并发布后才生效。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
